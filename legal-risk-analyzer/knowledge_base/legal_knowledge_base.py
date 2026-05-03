import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument

from config import Config
from database import LegalDocumentStore


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    content: str
    metadata: Dict[str, Any]
    chunks: List[str]


@dataclass
class RiskAnalysisResult:
    """风险分析结果"""
    risk_type: str
    description: str
    severity: str
    legal_basis: str
    location: str
    suggestions: str


class LegalKnowledgeBase:
    """
    法律文件知识库模块
    基于 PowerRAG 理念构建，实现：
    - 多格式文档解析 (PDF、Word、PPT 等)
    - 智能文档分块
    - 高质量检索结果生成
    """

    def __init__(self):
        self.doc_store = LegalDocumentStore()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
        )

        self.legal_risk_patterns = {
            "合同主体风险": [
                r"甲方.*?未明确",
                r"乙方.*?资质不全",
                r"签约主体.*?不一致",
            ],
            "条款缺失风险": [
                r"未约定.*?违约责任",
                r"缺少.*?争议解决条款",
                r"未明确.*?付款条件",
            ],
            "法律合规风险": [
                r"违反.*?法律规定",
                r"不符合.*?法规要求",
                r"超越.*?经营范围",
            ],
            "权利义务不对等": [
                r"单方面.*?解除合同",
                r"仅.*?承担.*?责任",
                r"不承担.*?任何责任",
            ],
            "保密条款风险": [
                r"保密期限.*?不明确",
                r"保密范围.*?过宽",
                r"未约定.*?保密义务",
            ],
            "知识产权风险": [
                r"知识产权.*?归属不明",
                r"侵权.*?责任.*?未约定",
                r"许可使用.*?范围不清",
            ],
        }

    def parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文件"""
        reader = PdfReader(file_path)
        text_content = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                text_content.append(f"[第{page_num + 1}页]\n{text.strip()}")

        content = "\n\n".join(text_content)
        chunks = self.text_splitter.split_text(content)

        return ParsedDocument(
            content=content,
            metadata={
                "source_file": Path(file_path).name,
                "file_type": "pdf",
                "page_count": len(reader.pages),
                "created_at": datetime.now().isoformat(),
            },
            chunks=chunks,
        )

    def parse_docx(self, file_path: str) -> ParsedDocument:
        """解析 Word 文件"""
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        content = "\n\n".join(paragraphs)
        chunks = self.text_splitter.split_text(content)

        return ParsedDocument(
            content=content,
            metadata={
                "source_file": Path(file_path).name,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "created_at": datetime.now().isoformat(),
            },
            chunks=chunks,
        )

    def parse_txt(self, file_path: str) -> ParsedDocument:
        """解析文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()

        chunks = self.text_splitter.split_text(content)

        return ParsedDocument(
            content=content,
            metadata={
                "source_file": Path(file_path).name,
                "file_type": "txt",
                "created_at": datetime.now().isoformat(),
            },
            chunks=chunks,
        )

    def parse_document(self, file_path: str) -> ParsedDocument:
        """根据文件类型解析文档"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self.parse_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        elif suffix == ".txt":
            return self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def add_document(
        self,
        file_path: str,
        doc_type: str = "contract",
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> int:
        """
        添加文档到知识库

        Args:
            file_path: 文件路径
            doc_type: 文档类型
            title: 文档标题
            metadata: 额外元数据

        Returns:
            添加的文档块数量
        """
        parsed = self.parse_document(file_path)

        combined_metadata = {
            **parsed.metadata,
            "doc_type": doc_type,
            "title": title or Path(file_path).stem,
            **(metadata or {}),
        }

        return self.doc_store.add_text(
            text=parsed.content,
            doc_id=Path(file_path).stem,
            doc_type=doc_type,
            title=title,
            metadata=combined_metadata,
        )

    def add_text_content(
        self,
        content: str,
        doc_id: str,
        doc_type: str = "contract",
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> int:
        """
        直接添加文本内容到知识库

        Args:
            content: 文本内容
            doc_id: 文档 ID
            doc_type: 文档类型
            title: 文档标题
            metadata: 额外元数据

        Returns:
            添加的文档块数量
        """
        combined_metadata = {
            "doc_type": doc_type,
            "title": title or doc_id,
            "created_at": datetime.now().isoformat(),
            **(metadata or {}),
        }

        return self.doc_store.add_text(
            text=content,
            doc_id=doc_id,
            doc_type=doc_type,
            title=title,
            metadata=combined_metadata,
        )

    def search(self, query: str, n_results: int = 5, doc_type: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        搜索知识库

        Args:
            query: 查询文本
            n_results: 返回结果数量
            doc_type: 过滤文档类型

        Returns:
            搜索结果列表
        """
        return self.doc_store.search(query, n_results, doc_type)

    def detect_risks_by_patterns(self, content: str) -> List[Dict[str, Any]]:
        """
        基于模式检测风险点

        Args:
            content: 文档内容

        Returns:
            检测到的风险点列表
        """
        detected_risks = []

        for risk_type, patterns in self.legal_risk_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, content)
                for match in matches:
                    start = max(0, match.start() - 50)
                    end = min(len(content), match.end() + 50)
                    context = content[start:end]

                    detected_risks.append({
                        "risk_type": risk_type,
                        "matched_text": match.group(),
                        "context": context,
                        "position": match.start(),
                    })

        return detected_risks

    def get_relevant_laws(self, risk_type: str) -> List[Dict[str, Any]]:
        """
        获取相关法律知识

        Args:
            risk_type: 风险类型

        Returns:
            相关法律知识列表
        """
        return self.doc_store.search_risk_knowledge(risk_type, n_results=3)

    def add_risk_knowledge(
        self,
        risk_type: str,
        description: str,
        legal_basis: str,
        severity: str = "medium",
        suggestions: Optional[str] = None,
    ):
        """
        添加风险知识

        Args:
            risk_type: 风险类型
            description: 风险描述
            legal_basis: 法律依据
            severity: 严重程度
            suggestions: 建议措施
        """
        self.doc_store.add_risk_knowledge(
            risk_type=risk_type,
            description=description,
            legal_basis=legal_basis,
            severity=severity,
            suggestions=suggestions,
        )

    def get_stats(self) -> Dict[str, Any]:
        """获取知识库统计信息"""
        return self.doc_store.get_stats()
